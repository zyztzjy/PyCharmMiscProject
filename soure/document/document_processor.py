# soure/document_processor/document_processor.py
import os
import fitz  # PyMuPDF
import pandas as pd
from typing import List, Dict, Optional, Any
import hashlib
import tempfile
from datetime import datetime
import re
from docx import Document  # 处理Word文档
import xlrd  # 处理旧版Excel
from openpyxl import load_workbook  # 处理新版Excel


class DocumentProcessor:
    """多格式文档处理器（支持PDF、Word、Excel）"""

    def __init__(self, config: dict = None):
        self.config = config or {}
        self.max_chunk_size = self.config.get('document', {}).get('max_chunk_size', 1000)
        self.overlap_size = self.config.get('document', {}).get('overlap_size', 200)

    def extract_text_from_document(self, file_path: str) -> List[Dict[str, Any]]:
        """从文档提取文本并分块，支持多种格式"""
        try:
            if not os.path.exists(file_path):
                print(f"文档文件不存在: {file_path}")
                return []

            # 根据文件扩展名选择处理方法
            file_ext = os.path.splitext(file_path)[1].lower()

            if file_ext == '.pdf':
                return self.extract_text_from_pdf(file_path)
            elif file_ext in ['.docx', '.doc']:
                return self.extract_text_from_word(file_path)
            elif file_ext in ['.xlsx', '.xls']:
                return self.extract_text_from_excel(file_path)
            else:
                print(f"不支持的文件格式: {file_ext}")
                return []

        except Exception as e:
            print(f"提取文档文本失败: {e}")
            return []

    def extract_text_from_pdf(self, pdf_path: str) -> List[Dict[str, Any]]:
        """从PDF提取文本并分块"""
        try:
            if not os.path.exists(pdf_path):
                print(f"PDF文件不存在: {pdf_path}")
                return []

            doc = fitz.open(pdf_path)
            full_text = ""

            # 从文件路径中提取企业名称
            file_name = os.path.basename(pdf_path)
            file_dir = os.path.dirname(pdf_path)

            # 尝试从目录名提取企业名称（如"撤否企业/欣强电子"）
            company_name = self._extract_company_name_from_path(file_dir, file_name)

            metadata = {
                "source": file_name,  # 文件名
                "file_path": pdf_path,  # 完整路径
                "file_name": file_name,  # 文件名
                "file_dir": file_dir,  # 目录路径
                "company_name": company_name,  # 新增：企业名称
                "file_size": os.path.getsize(pdf_path),
                "page_count": len(doc),
                "extraction_time": datetime.now().isoformat(),
                "document_type": self._detect_document_type(pdf_path)
            }

            # 提取所有页面文本
            for page_num in range(len(doc)):
                page = doc[page_num]
                text = page.get_text()
                if text.strip():
                    full_text += f"第{page_num + 1}页:\n{text}\n\n"

            doc.close()

            if not full_text.strip():
                print(f"PDF文件没有提取到文本: {pdf_path}")
                return []

            # 对文本进行分块
            chunks = self._chunk_text(full_text, metadata)

            print(f"从 {pdf_path} 提取了 {len(chunks)} 个文本块，企业名称: {company_name}")
            return chunks

        except Exception as e:
            print(f"提取PDF文本失败: {e}")
            return []

    def extract_text_from_word(self, doc_path: str) -> List[Dict[str, Any]]:
        """从Word文档提取文本并分块"""
        try:
            if not os.path.exists(doc_path):
                print(f"Word文档文件不存在: {doc_path}")
                return []

            # 从文件路径中提取企业名称
            file_name = os.path.basename(doc_path)
            file_dir = os.path.dirname(doc_path)

            # 尝试从目录名提取企业名称（如"撤否企业/欣强电子"）
            company_name = self._extract_company_name_from_path(file_dir, file_name)

            metadata = {
                "source": file_name,
                "file_path": doc_path,
                "file_name": file_name,
                "file_dir": file_dir,
                "company_name": company_name,
                "file_size": os.path.getsize(doc_path),
                "page_count": None,  # Word文档暂不计算页数
                "extraction_time": datetime.now().isoformat(),
                "document_type": self._detect_document_type(doc_path)
            }

            full_text = ""
            
            # 根据文件扩展名选择处理方式
            file_ext = os.path.splitext(doc_path)[1].lower()
            if file_ext == '.docx':
                # 使用python-docx处理.docx文件
                doc = Document(doc_path)
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        full_text += paragraph.text + "\n"
                
                # 处理表格内容
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                full_text += cell.text + "\n"
            elif file_ext == '.doc':
                # .doc格式需要特殊处理，这里简化处理
                print(f"警告: .doc格式可能不完全支持，请使用.docx格式以获得最佳效果: {doc_path}")
                # 对于旧版Word文档，我们暂时返回空结果，实际项目中可能需要使用其他库
                return []

            if not full_text.strip():
                print(f"Word文档没有提取到文本: {doc_path}")
                return []

            # 对文本进行分块
            chunks = self._chunk_text(full_text, metadata)

            print(f"从 {doc_path} 提取了 {len(chunks)} 个文本块，企业名称: {company_name}")
            return chunks

        except Exception as e:
            print(f"提取Word文档文本失败: {e}")
            return []

    def extract_text_from_excel(self, excel_path: str) -> List[Dict[str, Any]]:
        """从Excel文档提取文本并分块"""
        try:
            if not os.path.exists(excel_path):
                print(f"Excel文件不存在: {excel_path}")
                return []

            # 从文件路径中提取企业名称
            file_name = os.path.basename(excel_path)
            file_dir = os.path.dirname(excel_path)

            # 尝试从目录名提取企业名称（如"撤否企业/欣强电子"）
            company_name = self._extract_company_name_from_path(file_dir, file_name)

            metadata = {
                "source": file_name,
                "file_path": excel_path,
                "file_name": file_name,
                "file_dir": file_dir,
                "company_name": company_name,
                "file_size": os.path.getsize(excel_path),
                "page_count": None,  # Excel使用工作表数量
                "extraction_time": datetime.now().isoformat(),
                "document_type": self._detect_document_type(excel_path)
            }

            full_text = ""
            
            # 根据文件扩展名选择处理方式
            file_ext = os.path.splitext(excel_path)[1].lower()
            if file_ext == '.xlsx':
                # 使用openpyxl处理.xlsx文件
                workbook = load_workbook(excel_path, read_only=True)
                sheet_names = workbook.sheetnames
                
                for sheet_name in sheet_names:
                    worksheet = workbook[sheet_name]
                    
                    # 添加工作表名称
                    full_text += f"工作表: {sheet_name}\n"
                    
                    # 读取单元格数据
                    for row in worksheet.iter_rows(values_only=True):
                        for cell_value in row:
                            if cell_value is not None:
                                cell_text = str(cell_value).strip()
                                if cell_text:
                                    full_text += cell_text + "\n"
                    
                    full_text += "\n"  # 工作表之间添加分隔符
                
                workbook.close()
            elif file_ext == '.xls':
                # 使用xlrd处理.xls文件
                workbook = xlrd.open_workbook(excel_path)
                sheet_names = workbook.sheet_names()
                
                for sheet_name in sheet_names:
                    worksheet = workbook.sheet_by_name(sheet_name)
                    
                    # 添加工作表名称
                    full_text += f"工作表: {sheet_name}\n"
                    
                    # 读取单元格数据
                    for row_idx in range(worksheet.nrows):
                        for col_idx in range(worksheet.ncols):
                            cell_value = worksheet.cell_value(row_idx, col_idx)
                            if cell_value is not None:
                                cell_text = str(cell_value).strip()
                                if cell_text:
                                    full_text += cell_text + "\n"
                    
                    full_text += "\n"  # 工作表之间添加分隔符

            if not full_text.strip():
                print(f"Excel文件没有提取到文本: {excel_path}")
                return []

            # 对文本进行分块
            chunks = self._chunk_text(full_text, metadata)

            print(f"从 {excel_path} 提取了 {len(chunks)} 个文本块，企业名称: {company_name}")
            return chunks

        except Exception as e:
            print(f"提取Excel文档文本失败: {e}")
            return []

    def _extract_company_name_from_path(self, file_dir: str, file_name: str) -> str:
        """从文件路径中提取企业名称"""
        try:
            # 方法1：从文件名中提取（移除扩展名和常见词汇）
            base_name = os.path.splitext(file_name)[0]

            # 移除常见的文件描述词汇
            common_terms = [
                "股份有限公司", "有限公司", "公司",
                "招股说明书", "招股书", "招股",
                "审计报告", "审计", "报告",
                "法律意见书", "法律意见", "意见书",
                "发行保荐书", "发行保荐", "保荐书",
                "上市保荐书", "上市保荐",
                "财务报表", "财务报告", "财务",
                "年度报告", "年报", "报告",
                "资产负债表", "利润表", "现金流量表",
                "损益表", "权益变动表", "附注"
            ]

            # 尝试从文件名中提取
            for term in common_terms:
                if term in base_name:
                    # 保留企业名称部分
                    company_part = base_name.split(term)[0].strip()
                    if company_part:  # 如果分割后有内容
                        return company_part

            # 方法2：从目录名中提取
            dir_name = os.path.basename(file_dir)
            if dir_name and dir_name != ".":
                # 移除目录中的分类词汇
                category_terms = ["撤否企业", "长期辅导企业", "新三板企业",
                                  "供应链分析", "财务报告", "舆情报告", "行业分析",
                                  "Excel", "Word", "PDF"]
                for term in category_terms:
                    dir_name = dir_name.replace(term, "").strip()
                if dir_name:
                    return dir_name

            # 方法3：返回清理后的文件名
            return base_name

        except Exception as e:
            print(f"提取企业名称失败: {e}")
            return file_name  # 失败时返回原始文件名

    def _detect_document_type(self, file_path: str) -> str:
        """检测文档类型"""
        filename = os.path.basename(file_path).lower()

        if any(keyword in filename for keyword in ["财务", "年报", "审计", "报表", "finance", "资产负债", "利润", "现金流量"]):
            return "财务报告"
        elif any(keyword in filename for keyword in ["研报", "分析", "研究", "industry", "调研"]):
            return "研究报告"
        elif any(keyword in filename for keyword in ["招股", "ipo", "prospectus"]):
            return "招股说明书"
        elif any(keyword in filename for keyword in ["新闻", "舆情", "报道", "news"]):
            return "新闻舆情"
        elif any(keyword in filename for keyword in ["专利", "技术", "patent"]):
            return "技术专利"
        elif any(keyword in filename for keyword in ["报告", "document"]):
            return "报告文档"
        else:
            # 根据文件扩展名判断
            ext = os.path.splitext(filename)[1].lower()
            if ext == '.pdf':
                return "PDF文档"
            elif ext in ['.doc', '.docx']:
                return "Word文档"
            elif ext in ['.xls', '.xlsx']:
                return "Excel文档"
            else:
                return "其他文档"

    def _chunk_text(self, text: str, metadata: dict) -> List[Dict[str, Any]]:
        """将长文本分块"""
        chunks = []

        # 清理文本
        text = text.replace('\r\n', '\n').replace('\r', '\n')

        # 按段落分割
        paragraphs = re.split(r'\n\s*\n+', text)
        paragraphs = [p.strip() for p in paragraphs if p.strip()]

        if not paragraphs:
            # 如果没有明显段落，尝试按句子分割
            sentences = re.split(r'(?<=[。！？.!?])\s+', text)
            paragraphs = [s.strip() for s in sentences if s.strip()]

        current_chunk = ""
        current_length = 0

        for para in paragraphs:
            para_length = len(para)

            # 如果当前段落加上新段落超过限制，保存当前块并开始新块
            if current_length + para_length > self.max_chunk_size and current_chunk:
                # 保存当前块
                chunk_dict = self._create_chunk_dict(current_chunk, metadata)
                chunks.append(chunk_dict)

                # 开始新块（带重叠）
                if self.overlap_size > 0 and len(current_chunk) > self.overlap_size:
                    # 取最后overlap_size个字符作为重叠
                    overlap_text = current_chunk[-self.overlap_size:]
                    current_chunk = overlap_text + "\n\n" + para
                else:
                    current_chunk = para
                current_length = len(current_chunk)
            else:
                # 添加到当前块
                if current_chunk:
                    current_chunk += "\n\n" + para
                else:
                    current_chunk = para
                current_length = len(current_chunk)

            # 如果单个段落超过最大块大小，需要强制分割
            if para_length > self.max_chunk_size:
                # 先保存当前块
                if current_chunk and current_chunk != para:
                    chunk_dict = self._create_chunk_dict(current_chunk, metadata)
                    chunks.append(chunk_dict)

                # 分割超大段落
                sub_chunks = self._split_large_paragraph(para)
                for sub_chunk in sub_chunks:
                    chunk_dict = self._create_chunk_dict(sub_chunk, metadata)
                    chunks.append(chunk_dict)

                current_chunk = ""
                current_length = 0

        # 保存最后一个块
        if current_chunk:
            chunk_dict = self._create_chunk_dict(current_chunk, metadata)
            chunks.append(chunk_dict)

        return chunks

    def _create_chunk_dict(self, content: str, metadata: dict) -> Dict[str, Any]:
        """创建分块字典"""
        chunk_id = hashlib.md5(content.encode('utf-8')).hexdigest()[:16]

        return {
            "content": content,
            "metadata": {
                **metadata,
                "chunk_id": chunk_id,
                "chunk_size": len(content),
                "is_partial": True,
                "created_time": datetime.now().isoformat()
            }
        }

    def _split_large_paragraph(self, text: str) -> List[str]:
        """分割大段落文本"""
        # 按句子分割
        sentences = re.split(r'(?<=[。！？.!?])\s+', text)
        sentences = [s.strip() for s in sentences if s.strip()]

        chunks = []
        current_chunk = ""

        for sentence in sentences:
            if len(current_chunk) + len(sentence) <= self.max_chunk_size:
                if current_chunk:
                    current_chunk += " " + sentence
                else:
                    current_chunk = sentence
            else:
                if current_chunk:
                    chunks.append(current_chunk)
                current_chunk = sentence

        if current_chunk:
            chunks.append(current_chunk)

        return chunks

    def process_document_directory(self, directory_path: str) -> List[Dict[str, Any]]:
        """处理目录中的所有文档文件（支持PDF、Word、Excel）"""
        all_chunks = []

        if not os.path.exists(directory_path):
            print(f"目录不存在: {directory_path}")
            return all_chunks

        # 查找所有支持的文档文件
        doc_files = []
        for root, dirs, files in os.walk(directory_path):
            for file in files:
                file_lower = file.lower()
                if (file_lower.endswith('.pdf') or 
                    file_lower.endswith('.docx') or 
                    file_lower.endswith('.doc') or 
                    file_lower.endswith('.xlsx') or 
                    file_lower.endswith('.xls')):
                    doc_files.append(os.path.join(root, file))

        print(f"找到 {len(doc_files)} 个文档文件")

        # 处理每个文档文件
        for doc_file in doc_files:
            print(f"处理文件: {doc_file}")

            # 提取文本
            text_chunks = self.extract_text_from_document(doc_file)
            if text_chunks:
                all_chunks.extend(text_chunks)

        print(f"总共提取了 {len(all_chunks)} 个文本块")
        return all_chunks